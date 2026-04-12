"""DOCX parser — extracts paragraphs and tables from Word documents."""

import docx


def docxParser(file_path):
    """Extract text, tables, and document metadata from a .docx file."""
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    metadata = {
        "num_paragraphs": len(doc.paragraphs),
        "num_tables": len(doc.tables)
    }
    return {
        "text": text,
        "tables": tables,
        "metadata": metadata
    }